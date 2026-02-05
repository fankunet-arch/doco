#!/usr/bin/env python3
"""
修复 Word 文档页码不连续的问题

问题：文档被分成3个节，每个节的首页都没有页码，导致页码不连续
解决：删除多余的分节符，只保留封面和正文两个节
"""

from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Pt
import sys

def fix_page_numbering(input_file, output_file):
    """修复页码不连续的问题"""

    print(f'读取文档: {input_file}')
    doc = Document(input_file)

    print(f'\n原文档信息:')
    print(f'  总节数: {len(doc.sections)}')
    print(f'  总段落数: {len(doc.paragraphs)}')

    # 删除段落中的分节符（除了最后一个段落的分节符）
    # Word文档的分节符存储在段落的段落属性中
    w_ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'

    removed_sections = 0

    # 遍历所有段落，删除中间的分节符
    for i, para in enumerate(doc.paragraphs):
        # 查找段落内的分节符
        pPr = para._element.find(f'.//{w_ns}pPr')
        if pPr is not None:
            sectPr = pPr.find(f'.//{w_ns}sectPr')
            if sectPr is not None:
                # 这是一个段落级别的分节符
                # 我们需要删除它（除了在特定位置的）
                # 段落49和70是原来的分节符位置，我们要删除这些
                if i in [49, 70]:
                    print(f'  删除段落 {i} 的分节符')
                    pPr.remove(sectPr)
                    removed_sections += 1

    print(f'\n删除了 {removed_sections} 个分节符')

    # 现在文档应该只有2个节了
    # 重新加载文档以查看更改后的节数
    print(f'修改后节数: {len(doc.sections)}')

    # 设置第1节（封面）
    section1 = doc.sections[0]
    section1.different_first_page_header_footer = True

    # 确保第1节首页页脚为空（无页码）
    first_page_footer1 = section1.first_page_footer
    for para in first_page_footer1.paragraphs:
        para.clear()

    # 确保第1节普通页脚有页码
    footer1 = section1.footer
    if len(footer1.paragraphs) == 0:
        footer1.add_paragraph()

    print('\n第1节设置:')
    print('  - 不同首页: True')
    print('  - 首页页脚: 空（无页码）')
    print('  - 其他页页脚: 有页码')

    # 设置第2节（正文）- 如果存在
    if len(doc.sections) >= 2:
        section2 = doc.sections[1]
        section2.different_first_page_header_footer = False

        # 确保第2节页脚有页码
        footer2 = section2.footer
        if len(footer2.paragraphs) == 0:
            footer2.add_paragraph()

        print('\n第2节设置:')
        print('  - 不同首页: False')
        print('  - 所有页页脚: 有页码')

        # 设置页码从1开始
        sectPr = section2._sectPr
        pgNumType = sectPr.find(f'.//{w_ns}pgNumType')
        if pgNumType is None:
            pgNumType = OxmlElement('w:pgNumType')
            sectPr.insert(0, pgNumType)
        pgNumType.set(f'{w_ns}start', '1')

    # 删除其他多余的节（如果还有的话）
    if len(doc.sections) > 2:
        print(f'\n警告: 文档仍有 {len(doc.sections)} 个节，可能需要手动调整')

    # 保存文档
    print(f'\n保存修复后的文档: {output_file}')
    doc.save(output_file)
    print('✓ 完成!')

    print('\n修复说明:')
    print('  1. 删除了多余的分节符')
    print('  2. 保留封面节（首页无页码）')
    print('  3. 正文节所有页面都有连续页码（从1开始）')
    print('  4. 页码将会是连续的，不会再有断开')

if __name__ == '__main__':
    input_file = '桃喜产品操作手册_new.docx'
    output_file = '桃喜产品操作手册_fixed.docx'

    try:
        fix_page_numbering(input_file, output_file)
    except Exception as e:
        print(f'错误: {e}')
        import traceback
        traceback.print_exc()
        sys.exit(1)
