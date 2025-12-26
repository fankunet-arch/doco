#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
专业级 Markdown 转 Word 文档转换器
符合企业标准化文档格式要求
"""

import os
import re
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


class MarkdownToWordConverter:
    def __init__(self, output_dir='word'):
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)

    def setup_styles(self, doc):
        """设置文档样式"""
        # 设置默认字体
        doc.styles['Normal'].font.name = 'Microsoft YaHei'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        doc.styles['Normal'].font.size = Pt(11)
        doc.styles['Normal'].paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        doc.styles['Normal'].paragraph_format.line_spacing = 1.15
        doc.styles['Normal'].paragraph_format.space_after = Pt(8)

        # 一级标题
        h1 = doc.styles['Heading 1']
        h1.font.name = 'Microsoft YaHei'
        h1._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        h1.font.size = Pt(16)
        h1.font.bold = True
        h1.font.color.rgb = RGBColor(46, 116, 181)  # 深蓝色
        h1.paragraph_format.space_before = Pt(18)
        h1.paragraph_format.space_after = Pt(12)

        # 二级标题
        h2 = doc.styles['Heading 2']
        h2.font.name = 'Microsoft YaHei'
        h2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        h2.font.size = Pt(14)
        h2.font.bold = True
        h2.font.color.rgb = RGBColor(0, 0, 0)
        h2.paragraph_format.space_before = Pt(14)
        h2.paragraph_format.space_after = Pt(10)

        # 三级标题
        h3 = doc.styles['Heading 3']
        h3.font.name = 'Microsoft YaHei'
        h3._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        h3.font.size = Pt(12)
        h3.font.bold = True
        h3.paragraph_format.space_before = Pt(12)
        h3.paragraph_format.space_after = Pt(8)

        # 创建代码块样式
        try:
            code_style = doc.styles.add_style('CodeBlock', WD_STYLE_TYPE.PARAGRAPH)
        except:
            code_style = doc.styles['CodeBlock']

        code_style.font.name = 'Consolas'
        code_style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
        code_style.font.size = Pt(9)
        code_style.paragraph_format.left_indent = Cm(0.5)
        code_style.paragraph_format.space_before = Pt(6)
        code_style.paragraph_format.space_after = Pt(6)

        # 创建引用样式
        try:
            quote_style = doc.styles.add_style('Quote', WD_STYLE_TYPE.PARAGRAPH)
        except:
            quote_style = doc.styles['Quote']

        quote_style.font.name = 'Microsoft YaHei'
        quote_style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        quote_style.font.size = Pt(10)
        quote_style.font.italic = True
        quote_style.font.color.rgb = RGBColor(80, 80, 80)
        quote_style.paragraph_format.left_indent = Cm(1)

    def add_cover_page(self, doc, title):
        """添加封面页"""
        # 标题
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.paragraph_format.space_before = Cm(8)
        title_para.paragraph_format.space_after = Cm(2)

        title_run = title_para.add_run(title)
        title_run.font.name = 'Microsoft YaHei'
        title_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        title_run.font.size = Pt(28)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(46, 116, 181)

        # 装饰线
        line_para = doc.add_paragraph()
        line_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        line_run = line_para.add_run('─' * 30)
        line_run.font.color.rgb = RGBColor(46, 116, 181)
        line_para.paragraph_format.space_after = Cm(3)

        # 副标题
        subtitle_para = doc.add_paragraph()
        subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle_para.add_run('公司内部标准化文档')
        subtitle_run.font.name = 'Microsoft YaHei'
        subtitle_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        subtitle_run.font.size = Pt(14)
        subtitle_run.font.color.rgb = RGBColor(100, 100, 100)
        subtitle_para.paragraph_format.space_after = Cm(4)

        # 日期
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_para.add_run(f'生成日期：{datetime.now().strftime("%Y年%m月%d日")}')
        date_run.font.name = 'Microsoft YaHei'
        date_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        date_run.font.size = Pt(11)
        date_run.font.color.rgb = RGBColor(120, 120, 120)

        # 分页
        doc.add_page_break()

    def add_toc(self, doc):
        """添加目录"""
        toc_title = doc.add_paragraph()
        toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        toc_run = toc_title.add_run('目录')
        toc_run.font.name = 'Microsoft YaHei'
        toc_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        toc_run.font.size = Pt(16)
        toc_run.font.bold = True
        toc_title.paragraph_format.space_after = Pt(12)

        # 添加目录域
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar)

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
        run._r.append(instrText)

        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar)

        doc.add_page_break()

    def add_shading_to_paragraph(self, paragraph, color='F5F5F5'):
        """为段落添加背景色"""
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), color)
        paragraph._p.get_or_add_pPr().append(shading_elm)

    def add_border_to_paragraph(self, paragraph, color='2E74B5', width=12):
        """为段落添加左边框"""
        pPr = paragraph._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')

        left = OxmlElement('w:left')
        left.set(qn('w:val'), 'single')
        left.set(qn('w:sz'), str(width))
        left.set(qn('w:space'), '4')
        left.set(qn('w:color'), color)

        pBdr.append(left)
        pPr.append(pBdr)

    def parse_markdown_line(self, line):
        """解析单行 Markdown"""
        # 标题
        if line.startswith('# '):
            return ('h1', line[2:].strip())
        elif line.startswith('## '):
            return ('h2', line[3:].strip())
        elif line.startswith('### '):
            return ('h3', line[4:].strip())
        elif line.startswith('#### '):
            return ('h4', line[5:].strip())

        # 引用
        elif line.startswith('>'):
            return ('quote', line[1:].strip())

        # 列表
        elif re.match(r'^[\*\-\+]\s+', line):
            return ('bullet', re.sub(r'^[\*\-\+]\s+', '', line))
        elif re.match(r'^\d+[\.\)]\s+', line):
            return ('number', re.sub(r'^\d+[\.\)]\s+', '', line))

        # 代码块标记
        elif line.strip().startswith('```'):
            return ('code_fence', line.strip())

        # 分隔线
        elif re.match(r'^[\-\*]{3,}$', line.strip()):
            return ('hr', '')

        # 表格
        elif '|' in line and line.strip().startswith('|'):
            return ('table', line)

        # 普通文本
        else:
            return ('text', line)

    def apply_inline_formatting(self, paragraph, text):
        """应用内联格式（粗体、斜体、代码）"""
        # 处理粗体
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = paragraph.add_run(part[2:-2])
                run.font.bold = True
            else:
                # 处理行内代码
                code_parts = re.split(r'(`[^`]+`)', part)
                for code_part in code_parts:
                    if code_part.startswith('`') and code_part.endswith('`'):
                        run = paragraph.add_run(code_part[1:-1])
                        run.font.name = 'Consolas'
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor(199, 37, 78)
                    else:
                        paragraph.add_run(code_part)

    def parse_table(self, lines):
        """解析 Markdown 表格"""
        table_lines = []
        for line in lines:
            if '|' in line:
                cells = [cell.strip() for cell in line.split('|')]
                cells = [c for c in cells if c]  # 移除空单元格
                if cells and not all(re.match(r'^[\-:]+$', c) for c in cells):
                    table_lines.append(cells)
        return table_lines

    def convert_file(self, md_path):
        """转换单个 Markdown 文件"""
        print(f"正在处理: {os.path.basename(md_path)}")

        # 读取文件
        with open(md_path, 'r', encoding='utf-8') as f:
            content = f.read()

        # 提取标题
        title_match = re.match(r'^#\s+(.+)$', content, re.MULTILINE)
        doc_title = title_match.group(1) if title_match else os.path.basename(md_path)

        # 创建文档
        doc = Document()

        # 设置页面
        section = doc.sections[0]
        section.page_height = Cm(29.7)  # A4
        section.page_width = Cm(21)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

        # 设置样式
        self.setup_styles(doc)

        # 添加封面
        self.add_cover_page(doc, doc_title)

        # 添加目录
        self.add_toc(doc)

        # 解析内容
        lines = content.split('\n')
        i = 0
        in_code_block = False
        code_lines = []
        in_table = False
        table_lines = []

        while i < len(lines):
            line = lines[i]
            line_type, line_content = self.parse_markdown_line(line)

            # 代码块处理
            if line_type == 'code_fence':
                if not in_code_block:
                    in_code_block = True
                    code_lines = []
                else:
                    # 结束代码块
                    in_code_block = False
                    if code_lines:
                        code_para = doc.add_paragraph('\n'.join(code_lines))
                        code_para.style = 'CodeBlock'
                        self.add_shading_to_paragraph(code_para)
                    code_lines = []
                i += 1
                continue

            if in_code_block:
                code_lines.append(line)
                i += 1
                continue

            # 表格处理
            if line_type == 'table':
                if not in_table:
                    in_table = True
                    table_lines = []
                table_lines.append(line)
                i += 1
                # 检查下一行是否还是表格
                if i < len(lines):
                    next_type, _ = self.parse_markdown_line(lines[i])
                    if next_type != 'table':
                        # 表格结束
                        in_table = False
                        parsed_table = self.parse_table(table_lines)
                        if parsed_table and len(parsed_table) > 0:
                            # 创建表格
                            table = doc.add_table(rows=len(parsed_table), cols=len(parsed_table[0]))
                            table.style = 'Light Grid Accent 1'

                            for row_idx, row_data in enumerate(parsed_table):
                                for col_idx, cell_data in enumerate(row_data):
                                    if col_idx < len(table.rows[row_idx].cells):
                                        cell = table.rows[row_idx].cells[col_idx]
                                        cell.text = cell_data
                                        # 表头样式
                                        if row_idx == 0:
                                            for paragraph in cell.paragraphs:
                                                for run in paragraph.runs:
                                                    run.font.bold = True

                            doc.add_paragraph()  # 表格后添加空行
                        table_lines = []
                continue

            # 标题
            if line_type == 'h1':
                para = doc.add_heading(line_content, level=1)
            elif line_type == 'h2':
                para = doc.add_heading(line_content, level=2)
            elif line_type == 'h3':
                para = doc.add_heading(line_content, level=3)
            elif line_type == 'h4':
                para = doc.add_paragraph(line_content, style='Heading 3')

            # 引用
            elif line_type == 'quote':
                para = doc.add_paragraph()
                para.style = 'Quote'
                self.apply_inline_formatting(para, line_content)
                self.add_border_to_paragraph(para)

            # 列表
            elif line_type == 'bullet':
                para = doc.add_paragraph(style='List Bullet')
                self.apply_inline_formatting(para, line_content)
            elif line_type == 'number':
                para = doc.add_paragraph(style='List Number')
                self.apply_inline_formatting(para, line_content)

            # 分隔线
            elif line_type == 'hr':
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run('─' * 50)
                run.font.color.rgb = RGBColor(200, 200, 200)

            # 普通文本
            elif line_type == 'text' and line.strip():
                para = doc.add_paragraph()
                self.apply_inline_formatting(para, line)

            i += 1

        # 添加页脚（页码）
        section = doc.sections[0]
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        footer_para.text = '第 '

        # 添加页码域
        run = footer_para.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar1)

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        run._r.append(instrText)

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar2)

        footer_para.add_run(' 页')

        # 保存文档
        output_filename = os.path.splitext(os.path.basename(md_path))[0] + '.docx'
        output_path = os.path.join(self.output_dir, output_filename)
        doc.save(output_path)
        print(f"✓ 已生成: {output_filename}")

        return output_path


def main():
    """主函数"""
    docs_dir = 'docs'
    converter = MarkdownToWordConverter(output_dir='word')

    # 获取所有 Markdown 文件
    md_files = []
    for file in os.listdir(docs_dir):
        if file.endswith('.md'):
            md_files.append(os.path.join(docs_dir, file))

    # 排序
    md_files.sort()

    print(f"找到 {len(md_files)} 个 Markdown 文件")
    print("=" * 60)

    # 转换所有文件
    output_files = []
    for md_file in md_files:
        try:
            output_path = converter.convert_file(md_file)
            output_files.append(output_path)
        except Exception as e:
            print(f"✗ 处理失败 {os.path.basename(md_file)}: {str(e)}")

    print("=" * 60)
    print(f"转换完成！共生成 {len(output_files)} 个 Word 文档")
    print(f"输出目录: {os.path.abspath('word')}")

    return output_files


if __name__ == '__main__':
    main()
